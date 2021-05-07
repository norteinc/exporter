import os
import re
from datetime import datetime

import pythoncom
from PyQt5.QtCore import QThread, pyqtSignal
from docx import Document
from win32com import client as win32
from win32com.client import constants

from package.config import DIRNAME
from package.service import get_count_todo, getvedtype, parsemark, validved


class ValidateBills(QThread):
    """Методы для связи сигналов и слотов"""
    update_progress_bar = pyqtSignal(int)
    lock_ui = pyqtSignal(bool)
    set_progress_bar = pyqtSignal(int, int, bool)
    add_string_to_activity_log = pyqtSignal(str)
    add_string_to_error_log = pyqtSignal(str)

    def run(self):
        self.lock_ui.emit(True)
        self.add_string_to_activity_log.emit(
            "-------------------------------------------------------------")
        if get_count_todo() == 0:
            self.add_string_to_activity_log.emit("Нет ведомостей для проверки")
            return
        self.add_string_to_activity_log.emit("Начат процесс проверки ведомостей")
        self.add_string_to_activity_log.emit(
            "-------------------------------------------------------------")
        logfilename = "\\Список ошибок в ведомостях от " + datetime.strftime(datetime.now(),
                                                                             "%d_%m_%Y_%H_%M_%S") + ".txt"
        log_all = open(DIRNAME + logfilename, "w")
        files = os.listdir(DIRNAME)
        cnt_current = 0
        self.set_progress_bar.emit(0, len(files), True)
        for cur_file in files:
            if re.findall('~', cur_file):  # проверка на временные файлы
                continue
            if re.search("docx|doc|rtf|odt", cur_file):
                if not re.search("docx", cur_file):
                    pythoncom.CoInitialize()
                    self.add_string_to_activity_log.emit("Преобразую файл в формат *.docx: " + cur_file)
                    try:
                        word = win32.gencache.EnsureDispatch('Word.Application')
                    except AttributeError:  # очистка кэша для нормальной работы win32api
                        import sys
                        import shutil
                        MODULE_LIST = [m.__name__ for m in sys.modules.values()]
                        for module in MODULE_LIST:
                            if re.match(r'win32com\.gen_py\..+', module):
                                del sys.modules[module]
                        shutil.rmtree(os.path.join(os.environ.get('LOCALAPPDATA'), 'Temp', 'gen_py'))
                        from win32com import client  # запрос нового кэша для работы
                        word = win32.gencache.EnsureDispatch('Word.Application')
                    try:
                        doc = word.Documents.Open(DIRNAME + "//" + cur_file)
                    except:
                        resPath4 = DIRNAME + '\\Поврежденные ведомости'
                        if not (os.path.exists(resPath4)):
                            os.mkdir(resPath4)
                        os.rename(DIRNAME + "\\" + cur_file, resPath4 + "\\" + cur_file)
                        continue
                    doc.Activate()
                    new_file_abs = os.path.abspath(DIRNAME + "//" + cur_file)
                    new_file_abs = re.sub(r'\.\w+$', '.docx', new_file_abs)
                    word.ActiveDocument.SaveAs(new_file_abs, FileFormat=constants.wdFormatXMLDocument)
                    doc.Close(False)
                    os.remove(DIRNAME + "//" + cur_file)
                    cur_file = re.sub(r'\.\w+$', '.docx', cur_file)
            else:
                continue
            document = Document(DIRNAME + "\\" + cur_file)

            self.add_string_to_activity_log.emit("Обрабатываю ведомость: " + cur_file)

            cnt_current += 1
            self.update_progress_bar.emit(cnt_current)

            if len(document.tables) < 1:
                log_all.write("Файл " + cur_file + " не является ведомостью (нет таблицы).\n")
                self.add_string_to_error_log.emit("Файл " + cur_file + " не является ведомостью (нет таблицы).")
                resPath2 = DIRNAME + '\\Некорретные ведомости'
                if not (os.path.exists(resPath2)):
                    os.mkdir(resPath2)
                if not (os.path.exists(resPath2 + "\\" + cur_file)):
                    os.rename(DIRNAME + "\\" + cur_file, resPath2 + "\\" + cur_file)
                continue

            table = document.tables[0]
            discipline = table.rows[5].cells[4].text.strip()
            semester = table.rows[4].cells[17].text.strip()
            kaf = table.rows[6].cells[4].text.strip()
            # grname = table.rows[4].cells[27].text.strip()
            course = table.rows[4].cells[12].text.strip()
            faculty = table.rows[4].cells[7].text.strip()
            vedid = table.rows[7].cells[28].text.strip()
            vedid2 = table.rows[7].cells[6].text.strip()
            vedtype = table.rows[7].cells[0].text.strip()
            period = table.rows[7].cells[14].text.strip()
            dateexam = table.rows[len(table.rows) - 9].cells[7].text.strip()
            if re.search(r"ДОПОЛНИТЕЛЬНАЯ", table.rows[3].cells[17].text.strip()):
                peroidtype = "1"
            else:
                peroidtype = "0"
            vedtypeid = getvedtype(vedtype)
            if vedtypeid == 2:
                teacher = table.rows[6].cells[22].text.strip()
                grname = table.rows[4].cells[30].text.strip()
            else:
                teacher = table.rows[6].cells[19].text.strip()
                grname = table.rows[4].cells[27].text.strip()

            # блок проверки модификации ведомостей на добавление столбцов
            if (vedtypeid == 2 and len(table.columns) != 35) or (vedtypeid != 2 and len(table.columns) != 33):
                log_all.write(
                    "Файл " + cur_file + " изменен и не подлежит обработке. Требуется перезаполнить ведомость.\n")
                self.add_string_to_error_log.emit(
                    "Файл " + cur_file + " изменен и не подлежит обработке. Требуется перезаполнить ведомость.")
                resPath2 = DIRNAME + '\\Некорретные ведомости'
                if not (os.path.exists(resPath2)):
                    os.mkdir(resPath2)
                if not (os.path.exists(resPath2 + "\\" + cur_file)):
                    os.rename(DIRNAME + "\\" + cur_file, resPath2 + "\\" + cur_file)
                continue

            # блок проверки полей под таблицей
            if len(table.rows[len(table.rows) - 10].cells[0].text.strip()) > 0:
                log_all.write(
                    "В файле " + cur_file + " под таблицей имеется текст, или ведомость видоизменена. Удалите текст под таблицей, или перезаполните ведомость.\n")
                self.add_string_to_error_log.emit(
                    cur_file + " под таблицей имеется текст, или ведомость видоизменена. Удалите текст под таблицей, или перезаполните ведомость.")
                resPath2 = DIRNAME + '\\Некорретные ведомости'
                if not (os.path.exists(resPath2)):
                    os.mkdir(resPath2)
                if not (os.path.exists(resPath2 + "\\" + cur_file)):
                    os.rename(DIRNAME + "\\" + cur_file, resPath2 + "\\" + cur_file)
                continue
            # блок проверки технических полей
            if len(vedid) == 0 or len(vedid2) == 0 or not re.search("^[0-9]+", vedid) or not re.search("^[0-9]+",
                                                                                                       vedid2):
                log_all.write("В файле " + cur_file + " удалены технические поля. Требуется перезаполнить ведомость.\n")
                self.add_string_to_error_log.emit(
                    cur_file + " удалены технические поля. Требуется перезаполнить ведомость.")
                resPath2 = DIRNAME + '\\Некорретные ведомости'
                if not (os.path.exists(resPath2)):
                    os.mkdir(resPath2)
                if not (os.path.exists(resPath2 + "\\" + cur_file)):
                    os.rename(DIRNAME + "\\" + cur_file, resPath2 + "\\" + cur_file)
                continue

            # блок адекватности таблицы
            if len(discipline) == 0 or len(semester) == 0 or len(kaf) == 0 or len(grname) == 0 or len(
                    course) == 0 or len(
                faculty) == 0 or len(vedtype) == 0 or len(period) == 0:
                log_all.write("В файле " + cur_file + " нарушена сетка таблицы. Требуется перезаполнить ведомость.\n")
                self.add_string_to_error_log.emit(
                    cur_file + " - нарушена сетка таблицы. Требуется перезаполнить ведомость.")
                resPath2 = DIRNAME + '\\Некорретные ведомости'
                if not (os.path.exists(resPath2)):
                    os.mkdir(resPath2)
                if not (os.path.exists(resPath2 + "\\" + cur_file)):
                    os.rename(DIRNAME + "\\" + cur_file, resPath2 + "\\" + cur_file)
                continue

            # блок адекватности таблицы на зачетки
            fail = 0
            for i in range(10, len(table.rows) - 10):
                if vedtypeid == 2:
                    if len(table.rows[1].cells[21].text.strip()) > 10:
                        log_all.write(
                            "В файле " + cur_file + " повреждена таблица. Необходимо перезаполнить ведомость.\n")
                        self.add_string_to_error_log.emit(
                            cur_file + " - повреждена таблица. Необходимо перезаполнить ведомость.")
                        resPath2 = DIRNAME + '\\Некорретные ведомости'
                        if not (os.path.exists(resPath2)):
                            os.mkdir(resPath2)
                        if not (os.path.exists(resPath2 + "\\" + cur_file)):
                            os.rename(DIRNAME + "\\" + cur_file, resPath2 + "\\" + cur_file)
                        fail = 1
                        break
                else:
                    if len(table.rows[i].cells[15].text.strip()) > 10:
                        log_all.write(
                            "В файле " + cur_file + " повреждена таблица. Необходимо перезаполнить ведомость.\n")
                        self.add_string_to_error_log.emit(
                            cur_file + " - повреждена таблица. Необходимо перезаполнить ведомость.")
                        resPath2 = DIRNAME + '\\Некорретные ведомости'
                        if not (os.path.exists(resPath2)):
                            os.mkdir(resPath2)
                        if not (os.path.exists(resPath2 + "\\" + cur_file)):
                            os.rename(DIRNAME + "\\" + cur_file, resPath2 + "\\" + cur_file)
                        fail = 1
                        break
            if fail == 1:
                continue

            # если флаг продержится до конца ведомости - можно парсить
            check = 0

            # обработка 1 листа
            for i in range(10, len(table.rows) - 10):
                if vedtypeid == 2:  # Экзаменационные ведомости
                    if len(table.rows[i].cells[21].text.strip()) < 6:
                        log_all.write("В файле " + cur_file + " в строке №" + str(
                            i - 9) + " некорректный номер зачетной книжки\n")
                        self.add_string_to_error_log.emit(
                            cur_file + " в строке №" + str(i - 9) + " некорректный номер зачетной книжки.")
                        check = 1
                    markid1 = parsemark(table.rows[i].cells[9].text.upper(), vedtypeid, 1)
                    markid2 = parsemark(table.rows[i].cells[16].text.upper(), vedtypeid, 2)
                    if markid1 == -1 and markid2 == -1:
                        log_all.write("В файле " + cur_file + " в строке №" + str(
                            i - 9) + " недопустимое значение / пустое значение в полях 'допуск' и 'экзамен'\n")
                        self.add_string_to_error_log.emit(
                            cur_file + " в строке №" + str(
                                i - 9) + " недопустимое значение / пустое значение в полях 'допуск' и 'экзамен'")
                        check = 1
                        continue
                    if markid1 == -1:
                        log_all.write("В файле " + cur_file + " в строке №" + str(
                            i - 9) + " недопустимое значение / пустое значение в поле 'допуск'\n")
                        self.add_string_to_error_log.emit(
                            cur_file + " в строке №" + str(
                                i - 9) + " недопустимое значение / пустое значение в поле 'допуск'")
                        check = 1
                        continue
                    if markid2 == -1:
                        log_all.write("В файле " + cur_file + " в строке №" + str(
                            i - 9) + " недопустимое значение / пустое значение в поле 'экзамен'\n")
                        self.add_string_to_error_log.emit(
                            cur_file + " в строке №" + str(
                                i - 9) + " недопустимое значение / пустое значение в поле 'экзамен'")
                        check = 1
                        continue
                    if not validved(markid1, markid2):
                        if markid1 == 9:
                            log_all.write("В файле " + cur_file + " в строке №" + str(
                                i - 9) + " в поле 'допуск' значение 'не сдано', но есть оценка за экзамен или стоит 'не явился' (должно быть 'не допущен')\n")
                            self.add_string_to_error_log.emit(
                                cur_file + " в строке №" + str(
                                    i - 9) + " в поле 'допуск' значение 'не сдано', но есть оценка за экзамен или стоит 'не явился' (должно быть 'не допущен').")
                        else:
                            log_all.write("В файле " + cur_file + " в строке №" + str(
                                i - 9) + " в поле 'допуск' значение 'сдано', но за экзамен стоит 'не допущен'\n")
                            self.add_string_to_error_log.emit(
                                cur_file + " в строке №" + str(
                                    i - 9) + " в поле 'допуск' значение 'сдано', но за экзамен стоит 'не допущен'.")
                        check = 1
                else:
                    if len(table.rows[i].cells[15].text.strip()) < 6:
                        log_all.write("В файле " + cur_file + " в строке №" + str(
                            i - 9) + " некорректный номер зачетной книжки\n")
                        self.add_string_to_error_log.emit(
                            cur_file + " в строке №" + str(
                                i - 9) + " некорректный номер зачетной книжки.")
                        check = 1
                    markid = parsemark(table.rows[i].cells[10].text.upper(), vedtypeid, 0)
                    if markid == -1:
                        log_all.write("В файле " + cur_file + " в строке №" + str(
                            i - 9) + " недопустимое значение / пустое значение в полях 'оценка'\n")
                        self.add_string_to_error_log.emit(cur_file + " в строке №" + str(
                            i - 9) + " недопустимое значение / пустое значение в полях 'оценка'.")
                        check = 1
            # обработка 2 листа, если он есть
            if len(document.tables) == 2:
                table = document.tables[1]
                for i in range(8, len(table.rows) - 10):  # перебор строк в xml
                    string = table.rows[i]._tr.xml
                    num = re.findall(r"<w:t.*</w:t>", string)  # поиск подстрок
                    i = 0
                    # обработка подстроки
                    while i < len(num):
                        num[i] = re.sub(r'<w:t xml:space="preserve">', "", num[i])  # очистка от мусора
                        num[i] = re.sub(r'<w:t>', "", num[i])
                        num[i] = re.sub(r"</w:t>", "", num[i])
                        if len(num[i]) < 2:  # удаление фантомных строк
                            num.pop(i)
                            i = i - 1
                        i = i + 1
                    i = 0
                    # склейка подстроки
                    while i < len(num):
                        if re.search(r"НЕ", num[i].strip().upper()) and len(num[i]) < 5:  # склейка не
                            num[i] = num[i] + " " + num[i + 1]
                            num.pop(i + 1)
                        if re.search(r"\.", num[i].strip()) and len(num[i]) < 6:  # склейка инициалов
                            num[i - 1] = num[i - 1] + " " + num[i]
                            num.pop(i)
                            i = i - 1
                        i = i + 1
                    if vedtypeid != 2:
                        if len(num) > 4:
                            if re.findall(r"\d{3}[Б|С]\d{2}", num[4]) or re.findall(r"\d{4}", num[4]):
                                num[3] = num[4]
                    # разбор оценок
                    if vedtypeid == 2:  # Экзаменационные ведомости
                        if len(num) != 5:
                            log_all.write("В файле " + cur_file + " в строке №" + num[
                                0] + " пустое значение в полях 'допуск' и/или 'экзамен' или стоит Подпись преподавателя. Дополнительно попробуйте удалить оценку и ввести ее заново\n")
                            self.add_string_to_error_log.emit(cur_file + " в строке №" + num[
                                0] + " пустое значение в полях 'допуск' и/или 'экзамен' или стоит Подпись преподавателя. Дополнительно попробуйте удалить оценку и ввести ее заново.")
                            check = 1
                            continue
                        if len(num[4].upper()) < 6:
                            log_all.write("В файле " + cur_file + " в строке №" + num[
                                0] + " короткий номер зачетной книжки\n")
                            self.add_string_to_error_log.emit(cur_file + " в строке №" + num[
                                0] + " короткий номер зачетной книжки.")
                            check = 1
                        else:
                            if not re.findall(r"\d{1}", num[4]):
                                log_all.write("В файле " + cur_file + " в строке №" + num[
                                    0] + " не номер зачетной книжки\n")
                                self.add_string_to_error_log.emit(cur_file + " в строке №" + num[
                                    0] + " не номер зачетной книжки.")
                                check = 1
                        markid1 = parsemark(num[2].upper(), vedtypeid, 1)
                        markid2 = parsemark(num[3].upper(), vedtypeid, 2)
                        if markid1 == -1 and markid2 == -1:
                            log_all.write("В файле " + cur_file + " в строке №" + num[
                                0] + " недопустимое значение / пустое значение в полях 'допуск' и 'экзамен'\n")
                            self.add_string_to_error_log.emit(cur_file + " в строке №" + num[
                                0] + " недопустимое значение / пустое значение в полях 'допуск' и 'экзамен'")
                            check = 1
                            continue
                        if markid1 == -1:
                            log_all.write("В файле " + cur_file + " в строке №" + num[
                                0] + " недопустимое значение / пустое значение в поле 'допуск'\n")
                            self.add_string_to_error_log.emit(cur_file + " в строке №" + num[
                                0] + " недопустимое значение / пустое значение в поле 'допуск'")
                            check = 1
                            continue
                        if markid2 == -1:
                            log_all.write("В файле " + cur_file + " в строке №" + num[
                                0] + " недопустимое значение / пустое значение в поле 'экзамен'\n")
                            self.add_string_to_error_log.emit(cur_file + " в строке №" + num[
                                0] + " недопустимое значение / пустое значение в поле 'экзамен'")
                            check = 1
                            continue
                        if not validved(markid1, markid2):
                            if markid1 == 9:
                                log_all.write("В файле " + cur_file + " в строке №" + num[
                                    0] + " в поле 'допуск' значение 'не сдано', но есть оценка за экзамен или стоит 'не явился' (должно быть 'не допущен')\n")
                                self.add_string_to_error_log.emit(cur_file + " в строке №" + num[
                                    0] + " в поле 'допуск' значение 'не сдано', но есть оценка за экзамен или стоит 'не явился' (должно быть 'не допущен')")
                            else:
                                log_all.write("В файле " + cur_file + " в строке №" + num[
                                    0] + " в поле 'допуск' значение 'сдано', но за экзамен стоит 'не допущен'\n")
                                self.add_string_to_error_log.emit(cur_file + " в строке №" + num[
                                    0] + " в поле 'допуск' значение 'сдано', но за экзамен стоит 'не допущен'")
                            check = 1
                    else:
                        if len(num) < 4:
                            log_all.write("В файле " + cur_file + " в строке №" + num[
                                0] + " пустое значение в поле 'оценка'\n")
                            self.add_string_to_error_log.emit(cur_file + " в строке №" + num[
                                0] + " пустое значение в поле 'оценка'")
                            check = 1
                            continue
                        if len(num[3].upper()) < 6:
                            log_all.write("В файле " + cur_file + " в строке №" + num[
                                0] + " некорректный номер зачетной книжки\n")
                            self.add_string_to_error_log.emit(cur_file + " в строке №" + num[
                                0] + " некорректный номер зачетной книжки")
                            check = 1
                        else:
                            if not (re.findall(r"\d{3}[Б|С]\d{2}", num[3]) or re.findall(r"\d{4}", num[3])):
                                log_all.write("В файле " + cur_file + " в строке №" + num[
                                    0] + " внутренняя ошибка формата. Попробуйте удалить оценку и ввести ее заново\n")
                                self.add_string_to_error_log.emit(cur_file + " в строке №" + num[
                                    0] + " внутренняя ошибка формата. Попробуйте удалить оценку и ввести ее заново")
                                check = 1
                        markid = parsemark(num[2].upper(), vedtypeid, 0)
                        if markid == -1:
                            log_all.write("В файле " + cur_file + " в строке №" + num[
                                0] + " недопустимое значение / пустое значение в полях 'оценка' или стоит Подпись преподавателя. Дополнительно попробуйте удалить оценку и ввести ее заново\n")
                            self.add_string_to_error_log.emit(cur_file + " в строке №" + num[
                                0] + " недопустимое значение / пустое значение в полях 'оценка' или стоит Подпись преподавателя. Дополнительно попробуйте удалить оценку и ввести ее заново")
                            check = 1
            if check == 0:
                resPath1 = DIRNAME + '\\Готовы к выгрузке'
                if not (os.path.exists(resPath1)):
                    os.mkdir(resPath1)
                if not (os.path.exists(resPath1 + "\\" + cur_file)):
                    os.rename(DIRNAME + "\\" + cur_file, resPath1 + "\\" + cur_file)
                else:
                    log_all.write(
                        "Файл " + cur_file + " уже подготовлен к выгрузке в папке 'Готовы к выгрузке'. Удалите его из папки 'Готовы к выгрузке', чтобы выгрузить заново\n")
                    self.add_string_to_error_log.emit(
                        cur_file + " уже подготовлен к выгрузке в папке 'Готовы к выгрузке'. Удалите его из папки 'Готовы к выгрузке', чтобы выгрузить заново.")
        self.add_string_to_activity_log.emit(
            "-------------------------------------------------------------")
        self.set_progress_bar.emit(0, len(files), False)
        log_all.close()
        self.add_string_to_activity_log.emit("Процесс обработки ведомостей закончен")
        if os.path.getsize(DIRNAME + logfilename) == 0:
            os.remove(DIRNAME + logfilename)
        else:
            self.add_string_to_activity_log.emit(
                "Ошибки, возникшие при обработке ведомостей, можно найти в файле " + DIRNAME + logfilename)
        self.add_string_to_activity_log.emit(
            "-------------------------------------------------------------")
        return