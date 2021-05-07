import json
import os
import platform
import re
import zipfile
from datetime import datetime

import pythoncom
import xlsxwriter
from PyQt5.QtCore import QThread, pyqtSignal
from docx import Document
from win32com import client as win32

from package.config import DIRNAME, DATE_DISTANCE, ENV
from package.dict import slovar
from package.service import isConnected, send_mail, add_mark, getvedtype, getshortmarkbyid, parsemark


class ExportBills(QThread):
    """Методы для связи слотов и сигналов"""
    update_progress_bar = pyqtSignal(int)
    lock_ui = pyqtSignal(bool)
    set_progress_bar = pyqtSignal(int, int, bool)
    add_string_to_activity_log = pyqtSignal(str)
    add_string_to_error_log = pyqtSignal(str)

    def __init__(self, val):
        super().__init__()
        self.ph_date = val

    def run(self):
        self.lock_ui.emit(True)
        PC = platform.uname()
        SNNumber = datetime.strftime(datetime.now(), "%d_%m_%Y_%H_%M_%S") + PC.node + " " + PC.machine + str(
            os.urandom(10))
        intSNNumber = abs(hash(SNNumber)) % (10 ** 8)
        while intSNNumber < 9999999:
            SNNumber = datetime.strftime(datetime.now(), "%d_%m_%Y_%H_%M_%S") + PC.node + " " + PC.machine + str(
                os.urandom(10))
            intSNNumber = abs(hash(SNNumber)) % (10 ** 8)
        username = PC.node
        for key in slovar:
            username = username.replace(key, slovar[key])
        dirFrom = r"C:\Ведомости\Готовы к выгрузке"
        files = os.listdir(dirFrom)
        dirExported = r"C:\Ведомости\Выгруженные ведомости_" + str(intSNNumber)
        ListVed = []
        NameVed = []
        curNumber = 1
        duplicate = 0
        cnt_current = 0
        self.set_progress_bar.emit(0, len(files), True)
        self.update_progress_bar.emit(cnt_current)
        self.add_string_to_activity_log.emit(
            "-------------------------------------------------------------")
        self.add_string_to_activity_log.emit("Происходит проверка дубликатов.")
        for cur_file in files:
            if re.findall('~', cur_file):  # проверка на временные файлы
                continue
            if not re.search("docx", cur_file):
                continue
            document = Document(dirFrom + "\\" + cur_file)
            table = document.tables[0]
            curNumber += 1
            vedid = table.rows[7].cells[28].text.strip()
            check = 0
            i = 0
            while i < len(ListVed):
                if vedid == ListVed[i]:
                    check = 1
                    break
                i += 1
            if check == 0:
                ListVed.append(vedid)
                NameVed.append(cur_file)
            else:
                self.add_string_to_error_log.emit(
                    "Дубликат ведомости: <" + cur_file + "> ИД " + vedid + " и <" + NameVed[i] + "> ИД " + ListVed[i])
                duplicate += 1
            cnt_current += 1
            self.update_progress_bar.emit(cnt_current)
        self.set_progress_bar.emit(0, len(files), False)
        if duplicate == 0:
            self.add_string_to_activity_log.emit("Дубликатов не обранаружено!")
        else:
            self.add_string_to_activity_log.emit(
                "Обнаружены дубликаты ведомостей. Оставьте в выгрузке только одну ведомость.")
            return
        self.add_string_to_activity_log.emit(
            "-------------------------------------------------------------")
        exportfilename = "\\Data_Rating_" + datetime.strftime(datetime.now(), "%d_%m_%Y_%H_%M_%S_") + str(
            intSNNumber) + "_" + str(username) + ".xlsx"

        exportfilestatsname = "\\Stats_" + datetime.strftime(datetime.now(), "%d_%m_%Y_%H_%M_%S_") + str(
            intSNNumber) + "_" + str(username) + ".xlsx"
        wbStatsOut = xlsxwriter.Workbook(DIRNAME + exportfilestatsname)
        wsStatsOut = wbStatsOut.add_worksheet("DR")
        wbOut = xlsxwriter.Workbook(DIRNAME + exportfilename)
        wsOut = wbOut.add_worksheet("DR")
        exportstr = 1
        wsOut.set_column("A:A", 9)
        wsOut.set_column("B:G", 7)
        wsOut.set_column("H:H", 40)
        wsOut.set_column("I:I", 8)
        wsOut.set_column("J:J", 20)
        wsOut.set_column("K:O", 8)
        wsOut.set_column("P:P", 65)
        wsOut.set_column("Q:Q", 20)
        wsOut.set_column("R:R", 10)
        wsStatsOut.set_column("A:A", 9)
        wsStatsOut.set_column("B:B", 40)
        wsStatsOut.set_column("C:D", 8)
        wsStatsOut.set_column("E:E", 25)
        wsStatsOut.set_column("F:F", 7)
        wsStatsOut.set_column("G:G", 9)
        wsStatsOut.set_column("H:H", 30)
        wsStatsOut.set_column("I:S", 10)
        all_format = wbOut.add_format(dict(align='center', valign='vcenter', text_wrap=1, font_size="11"))
        bold_format = wbOut.add_format(dict(bold=1, align='center', valign='vcenter', text_wrap=1, font_size="11"))
        all_format_stats = wbStatsOut.add_format(dict(align='center', valign='vcenter', text_wrap=1, font_size="11"))
        bold_format_stats = wbStatsOut.add_format(
            dict(bold=1, align='center', valign='vcenter', text_wrap=1, font_size="11"))
        OutHeaders = ["SN_Pack", "Number_Str", "id1", "idDis", "SemestrNumber", "CourseNumber", "isAdditional",
                      "PeriodName",
                      "StudRecBook",
                      "StudentFIO", "ControlTypeCode", "RatingCode", "FacultetShortName", "GroupNumber",
                      "KafedraNumber", "DisciplineName", "Ekzamenator", "DateExam"]

        OutHeadersStats = ["Кафедра", "Дисциплина", "Форма аттестации", "Семестр", "Период", "Этап", "Группа",
                           "Преподаватель", "Всего студентов", "2 / Не зач. (кол-во)",
                           "Доп (кол-во)", "Зач (кол-во)", "3 / Зач-3 (кол-во)", "4 / Зач - 4 (кол-во)",
                           "5 / Зач - 5 (кол-во)",
                           "ИП", "Сдавших (%)", "Не сдавших (%)", "Комментарий"]
        wsOut.write_row("A1", OutHeaders, bold_format)
        wsStatsOut.write_row("A1", OutHeadersStats, bold_format_stats)
        vedcount = 0
        self.add_string_to_activity_log.emit("Уникальный идентификатор выгрузки: " + str(intSNNumber))
        self.add_string_to_activity_log.emit(
            "-------------------------------------------------------------")
        self.set_progress_bar.emit(0, len(files), True)
        self.update_progress_bar.emit(vedcount)
        for cur_file in files:
            if re.findall('~', cur_file):  # проверка на временные файлы
                continue
            if not re.search("docx", cur_file):
                continue
            document = Document(dirFrom + "\\" + cur_file)
            self.add_string_to_activity_log.emit("Обрабатываю ведомость: " + cur_file)
            dict_file = {}
            table = document.tables[0]
            discipline = table.rows[5].cells[4].text.strip()
            semester = table.rows[4].cells[17].text.strip()
            kaf = table.rows[6].cells[4].text.strip()
            course = table.rows[4].cells[12].text.strip()
            faculty = table.rows[4].cells[7].text.strip()
            vedid = table.rows[7].cells[28].text.strip()
            vedid2 = table.rows[7].cells[6].text.strip()
            vedtype = table.rows[7].cells[0].text.strip()
            period = table.rows[7].cells[14].text.strip()
            if re.search(r"ДОПОЛНИТЕЛЬНАЯ", table.rows[3].cells[17].text.strip()):
                peroidtype = "1"
            else:
                peroidtype = "0"
            vedtypeid = getvedtype(vedtype)
            if vedtypeid == 2:
                teacher = table.rows[6].cells[22].text.strip()
                if re.findall(r"Экзаменатор", teacher):
                    teacher = table.rows[6].cells[24].text.strip()
                grname = table.rows[4].cells[30].text.strip()
            else:
                teacher = table.rows[6].cells[19].text.strip()
                if re.findall(r"Экзаменатор", teacher):
                    teacher = table.rows[6].cells[23].text.strip()
                grname = table.rows[4].cells[27].text.strip()
            if len(teacher) < 3:
                teacher = "Иванов И.И."
            dict_file["order"] = []
            dateexam = table.rows[len(table.rows) - 9].cells[7].text.strip()  # получаем дату из ячейки
            if len(dateexam) < 8:  # если дата не дата (слишком короткая)
                dateexam = self.ph_date  # ставим заглушку
            else:
                if re.search(r"\d{2}\.\d{2}\.\d{4}", dateexam):  # пытаемся найти подстрочку дд.мм.гггг
                    dateexam = re.search(r"\d{2}\.\d{2}\.\d{4}", dateexam)
                    dateexam = dateexam[0]
                elif re.search(r"\d{2}\.\d{2}\.\d{2}", dateexam):  # пытаемся найти подстрочку дд.мм.гг
                    dateexam = re.search(r"\d{2}\.\d{2}\.\d{2}", dateexam)
                    dateexam = dateexam[0]
                    try:  # проверка на то, является ли это датой
                        temp_dateexam = datetime.strptime(dateexam, "%d.%m.%y")
                    except ValueError:
                        dateexam = self.ph_date
                    else:
                        dateexam = datetime.strftime(temp_dateexam, "%d.%m.%Y")  # приводим дату к формату дд.мм.гггг
            try:
                validate_year = datetime.strptime(dateexam,
                                                  "%d.%m.%Y")  # получаем дату для проверки на срок годности, заодно проверяем является ли это датой
            except ValueError:
                dateexam = self.ph_date
            else:
                if abs(validate_year - datetime.now()).days > DATE_DISTANCE:  # проверяем, не далеко ли от сегодняшнего дня
                    dateexam = self.ph_date
            # обработка 1 листа
            stats_counter = {'total': 0, 'dopusk': 0, 'zach': 0, 'five': 0, 'four': 0, 'three': 0, 'two': 0, 'ip': 0}
            for i in range(10, len(table.rows) - 10):
                if vedtypeid == 2:  # Экзаменационные ведомости
                    markid1 = parsemark(table.rows[i].cells[9].text.upper(), vedtypeid, 1)
                    markid2 = parsemark(table.rows[i].cells[16].text.upper(), vedtypeid, 2)
                    # формируем строку в нужном формате
                    export1 = [intSNNumber, exportstr, vedid, vedid2, semester, course, peroidtype, period,
                               table.rows[i].cells[21].text.strip(), table.rows[i].cells[1].text.strip(),
                               "КМ", getshortmarkbyid(markid1), faculty, grname, kaf, discipline,
                               teacher, dateexam]
                    export2 = [intSNNumber, exportstr + 1, vedid, vedid2, semester, course, peroidtype, period,
                               table.rows[i].cells[21].text.strip(), table.rows[i].cells[1].text.strip(),
                               vedtype, getshortmarkbyid(markid2), faculty, grname, kaf, discipline,
                               teacher, dateexam]
                    # добавляем в словарь для быстрой обработки
                    dict_file[table.rows[i].cells[21].text.strip()] = []
                    dict_file[table.rows[i].cells[21].text.strip()].append(
                        [table.rows[i].cells[1].text.strip(), "КМ", getshortmarkbyid(markid1)])
                    dict_file[table.rows[i].cells[21].text.strip()].append(
                        [table.rows[i].cells[1].text.strip(), vedtype, getshortmarkbyid(markid2)])
                    dict_file["order"].append(
                        [table.rows[i].cells[1].text.strip(), table.rows[i].cells[21].text.strip()])
                    add_mark(stats_counter, markid1)
                    add_mark(stats_counter, markid2)
                    exportstr += 2
                    wsOut.write_row("A" + str(exportstr - 1), export1, all_format)
                    wsOut.write_row("A" + str(exportstr), export2, all_format)
                else:
                    markid = parsemark(table.rows[i].cells[10].text.upper(), vedtypeid, 0)
                    # формируем строку в нужном формате
                    export = [intSNNumber, exportstr, vedid, vedid2, semester, course, peroidtype, period,
                              table.rows[i].cells[15].text.strip(), table.rows[i].cells[1].text.strip(),
                              vedtype, getshortmarkbyid(markid), faculty, grname, kaf, discipline,
                              teacher, dateexam]
                    # добавляем в словарь для быстрой обработки
                    dict_file[table.rows[i].cells[21].text.strip()] = [table.rows[i].cells[1].text.strip(), vedtype,
                                                                       getshortmarkbyid(markid)]
                    dict_file["order"].append(
                        [table.rows[i].cells[1].text.strip(), table.rows[i].cells[21].text.strip()])
                    add_mark(stats_counter, markid)
                    exportstr += 1
                    wsOut.write_row("A" + str(exportstr), export, all_format)
            # обработка 2 листа, если он есть
            if len(document.tables) == 2:
                table = document.tables[1]
                for i in range(8, len(table.rows) - 10):  # перебор строк в xml
                    string = table.rows[i]._tr.xml
                    num = re.findall(r"<w:t.*</w:t>", string)  # поиск подстрок
                    i = 0
                    # обработка подстроки
                    while i < len(num):
                        num[i] = re.sub(r'<w:t xml:space="preserve">', "", num[i])  # очистка от мусора
                        num[i] = re.sub(r'<w:t>', "", num[i])
                        num[i] = re.sub(r"</w:t>", "", num[i])
                        if len(num[i]) < 2:  # удаление фантомных строк
                            num.pop(i)
                            i = i - 1
                        i = i + 1
                    i = 0
                    # склейка подстроки
                    while i < len(num):
                        if re.search(r"НЕ", num[i].strip().upper()) and len(num[i]) < 5:  # склейка не
                            num[i] = num[i] + " " + num[i + 1]
                            num.pop(i + 1)
                        if re.search(r"\.", num[i].strip()) and len(num[i]) < 6:  # склейка инициалов
                            num[i - 1] = num[i - 1] + " " + num[i]
                            num.pop(i)
                            i = i - 1
                        i = i + 1
                    if vedtypeid != 2:
                        if len(num) > 4:
                            if re.findall(r"\d{3}[Б|С]\d{2}", num[3]) or re.findall(r"\d{4}", num[3]):
                                num[3] = num[4]
                    num[1] = re.sub(r'  ', " ", num[1])
                    # разбор оценок
                    if vedtypeid == 2:  # Экзаменационные ведомости
                        markid1 = parsemark(num[2].upper(), vedtypeid, 1)
                        markid2 = parsemark(num[3].upper(), vedtypeid, 2)
                        #  формируем строку в нужном формате
                        export1 = [intSNNumber, exportstr, vedid, vedid2, semester, course, peroidtype, period,
                                   num[4].upper(), num[1].strip(),
                                   "КМ", getshortmarkbyid(markid1), faculty, grname, kaf, discipline,
                                   teacher, dateexam]
                        export2 = [intSNNumber, exportstr + 1, vedid, vedid2, semester, course, peroidtype, period,
                                   num[4].upper(), num[1].strip(),
                                   vedtype, getshortmarkbyid(markid2), faculty, grname, kaf, discipline,
                                   teacher, dateexam]
                        # добавляем в словарь для быстрой обработки
                        dict_file[num[4].upper()] = []
                        dict_file[num[4].upper()].append(
                            [num[1].strip(), "КМ", getshortmarkbyid(markid1)])
                        dict_file[num[4].upper()].append(
                            [num[1].strip(), vedtype, getshortmarkbyid(markid2)])
                        dict_file["order"].append([num[1].strip(), num[4].upper()])
                        add_mark(stats_counter, markid1)
                        add_mark(stats_counter, markid2)
                        exportstr += 2
                        wsOut.write_row("A" + str(exportstr - 1), export1, all_format)
                        wsOut.write_row("A" + str(exportstr), export2, all_format)
                    else:
                        markid = parsemark(num[2].upper(), vedtypeid, 0)
                        # формируем строку в нужном формате
                        export = [intSNNumber, exportstr, vedid, vedid2, semester, course, peroidtype, period,
                                  num[3].upper(), num[1].strip(),
                                  vedtype, getshortmarkbyid(markid), faculty, grname, kaf, discipline,
                                  teacher, dateexam]
                        # добавляем в словарь для быстрой обработки
                        dict_file[num[3].upper()] = [num[1].strip(), vedtype, getshortmarkbyid(markid)]
                        dict_file["order"].append([num[1].strip(), num[3].upper()])
                        add_mark(stats_counter, markid)
                        exportstr += 1
                        wsOut.write_row("A" + str(exportstr), export, all_format)
            vedcount += 1
            period = re.sub("[\/]", "_", period)  # убираем / из названия папки
            discipline = re.sub("[\.\:\;\/\\\|\?\*\"\<\>]", ".",
                                discipline)  # убираем все запретные символы из названия файла
            dict_file["meta"] = [grname, discipline, vedtype, semester, course, kaf, teacher, period]
            self.save_dict(dict_file)
            positive_results = round((((stats_counter['zach'] + stats_counter['three'] + stats_counter['four'] +
                                        stats_counter['five'] + stats_counter['ip']) * 100) / stats_counter['total']),
                                     2)
            negative_results = round(((stats_counter['two'] * 100) / stats_counter['total']), 2)
            comment = ""
            if positive_results < 25:
                comment = "<25% сдали"
            if positive_results > 90:
                comment = ">90% сдали"
            if positive_results == 100:
                comment = "100% ведомость"
            stats = [kaf, discipline, vedtype, semester, period, peroidtype, grname, teacher,
                     str(stats_counter['total']), str(stats_counter['two']),
                     str(stats_counter['dopusk']), str(stats_counter['zach']), str(stats_counter['three']),
                     str(stats_counter['four']),
                     str(stats_counter['five']), str(stats_counter['ip']),
                     str(positive_results),
                     str(negative_results), comment]
            wsStatsOut.write_row("A" + str(vedcount + 1), stats, all_format_stats)
            self.update_progress_bar.emit(vedcount)
            if not (os.path.exists(dirExported)):
                os.mkdir(dirExported)
            if not (os.path.exists(dirExported + "\\" + cur_file)):
                os.rename(dirFrom + "\\" + cur_file, dirExported + "\\" + cur_file)
            else:
                os.rename(dirFrom + "\\" + cur_file,
                          dirExported + "\\" + datetime.strftime(datetime.now(), "%d_%m_%Y_%H_%M_%S ") + cur_file)
        wsOut.name = "DR_" + str(intSNNumber) + "_V" + str(vedcount) + "_ST" + str(exportstr - 1)
        wbOut.close()
        wbStatsOut.close()
        if exportstr != 1:
            pythoncom.CoInitialize()
            excel = win32.gencache.EnsureDispatch('Excel.Application')
            wb = excel.Workbooks.Open(DIRNAME + exportfilename)
            new_file_abs = os.path.abspath(DIRNAME + exportfilename)
            new_file_abs = re.sub(r'\.\w+$', '_VED_' + str(vedcount) + '.xls', new_file_abs)
            wb.SaveAs(new_file_abs, FileFormat=56)
            excel.Application.Quit()
            os.remove(DIRNAME + exportfilename)
            exportfilename = new_file_abs
            z = zipfile.ZipFile(DIRNAME + "\\Выгруженные ведомости_" + str(intSNNumber) + ".zip",
                                'w')  # Создание нового архива
            files = os.listdir(dirExported)
            for file in files:  # Список всех файлов и папок в директории folder
                z.write(dirExported + "\\" + file)  # Создание относительных путей и запись файлов в архив
            z.close()
            self.add_string_to_activity_log.emit(
                "-------------------------------------------------------------")
            self.add_string_to_activity_log.emit(
                "Пожалуйста, подождите, идет создание архива.")
            if ENV == "RELEASE":
                if isConnected():
                    send_mail(DIRNAME + exportfilestatsname)
            self.add_string_to_activity_log.emit(
                "-------------------------------------------------------------")
            self.add_string_to_activity_log.emit(
                "Обработка успешно завершена.")
            self.add_string_to_activity_log.emit(
                "Все выгруженные ведомости перемещены в папку 'Выгруженные ведомости_" + str(
                    intSNNumber) + "'")
            self.add_string_to_activity_log.emit("Файл для загрузки в ИС УМУ - " + exportfilename)
            self.add_string_to_activity_log.emit(
                "Архив с ведомостями для загрузки в ИС УМУ - " + "Выгруженные ведомости_" + str(
                    intSNNumber) + ".zip")
            self.add_string_to_activity_log.emit(
                "-------------------------------------------------------------")
        else:
            self.add_string_to_activity_log.emit("Ведомостей для выгрузки не найдено")
            os.remove(DIRNAME + exportfilename)
        self.set_progress_bar.emit(0, len(files), False)
        return

    def save_dict(self, dict_file):
        """Сохранение словаря для дальнейшей работы с выгрузкой"""
        gr_name = dict_file["meta"][0]
        discipline = dict_file["meta"][1]
        kaf = dict_file["meta"][5]
        period = dict_file["meta"][7]
        sys_dir = DIRNAME + "\\" + ".sys" + "\\" + period + "\\" + gr_name
        if not os.path.exists(sys_dir):
            try:
                os.makedirs(sys_dir)
            except:
                self.add_string_to_error_log("Не удалось создать папку для сохранения данных")
                return
        try:
            a_file = open(sys_dir + "\\" + gr_name + "_" + discipline + "_" + kaf + ".json", "w")
        except:
            self.add_string_to_error_log("Не удалось сохранить данные в словарь")
            return
        json.dump(dict_file, a_file, ensure_ascii=False)
